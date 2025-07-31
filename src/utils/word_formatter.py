import os
import tempfile
import re
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class ProfessionalWordFormatter:
    """Formatador profissional de documentos Word para ETP"""
    
    def __init__(self):
        self.blue_color = RGBColor(31, 78, 121)  # Azul escuro #1f4e79
        self.white_color = RGBColor(255, 255, 255)  # Branco
        
    def create_professional_document(self, content: str, session_data: Dict = None) -> str:
        """Cria documento Word com formatação profissional"""
        try:
            doc = Document()
            
            # Configurar margens e layout
            self._configure_page_layout(doc)
            
            # Criar estilos personalizados
            self._create_custom_styles(doc)
            
            # Adicionar cabeçalho
            self._add_header(doc)
            
            # Adicionar título principal
            self._add_main_title(doc)
            
            # Processar e adicionar conteúdo
            self._process_and_add_content(doc, content)
            
            # Adicionar rodapé
            self._add_footer(doc)
            
            # Adicionar seção de assinatura
            self._add_signature_section(doc, session_data)
            
            # Salvar documento
            doc_path = self._save_document(doc)
            
            return doc_path
            
        except Exception as e:
            raise Exception(f"Erro ao criar documento Word: {str(e)}")
    
    def _configure_page_layout(self, doc: Document):
        """Configura layout da página conforme especificações"""
        sections = doc.sections
        for section in sections:
            # Margens: superior/inferior 2,5cm, esquerda/direita 3cm
            section.top_margin = Inches(0.98)    # 2,5cm
            section.bottom_margin = Inches(0.98) # 2,5cm
            section.left_margin = Inches(1.18)   # 3cm
            section.right_margin = Inches(1.18)  # 3cm
            
            # Configurar cabeçalho e rodapé
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)
    
    def _create_custom_styles(self, doc: Document):
        """Cria estilos personalizados para o documento"""
        styles = doc.styles
        
        # Estilo para títulos principais (com fundo azul)
        if 'Titulo Principal ETP' not in [s.name for s in styles]:
            title_style = styles.add_style('Titulo Principal ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            # Formatação do parágrafo
            title_format = title_style.paragraph_format
            title_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_format.space_before = Pt(12)
            title_format.space_after = Pt(12)
            title_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            # Formatação da fonte
            title_font = title_style.font
            title_font.name = 'Arial'
            title_font.size = Pt(14)
            title_font.bold = True
            title_font.color.rgb = self.white_color
        
        # Estilo para subtítulos
        if 'Subtitulo ETP' not in [s.name for s in styles]:
            subtitle_style = styles.add_style('Subtitulo ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            subtitle_format = subtitle_style.paragraph_format
            subtitle_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            subtitle_format.space_before = Pt(8)
            subtitle_format.space_after = Pt(6)
            subtitle_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            
            subtitle_font = subtitle_style.font
            subtitle_font.name = 'Arial'
            subtitle_font.size = Pt(12)
            subtitle_font.bold = True
            subtitle_font.color.rgb = RGBColor(0, 0, 0)
        
        # Estilo para corpo do texto
        if 'Corpo Texto ETP' not in [s.name for s in styles]:
            body_style = styles.add_style('Corpo Texto ETP', WD_STYLE_TYPE.PARAGRAPH)
            
            body_format = body_style.paragraph_format
            body_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_format.line_spacing = 1.5
            body_format.first_line_indent = Inches(0.49)  # 1,25cm
            body_format.space_after = Pt(6)
            
            body_font = body_style.font
            body_font.name = 'Arial'
            body_font.size = Pt(12)
            body_font.color.rgb = RGBColor(0, 0, 0)
    
    def _add_header(self, doc: Document):
        """Adiciona cabeçalho institucional"""
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        
        # Limpar parágrafo existente
        header_para.clear()
        
        # Adicionar texto do cabeçalho
        run = header_para.add_run("GOVERNO DO ESTADO")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        header_para.add_run("\n")
        
        run = header_para.add_run("SECRETARIA DE ADMINISTRAÇÃO")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        header_para.add_run("\n")
        
        run = header_para.add_run("ESTUDO TÉCNICO PRELIMINAR")
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar linha separadora
        header.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_main_title(self, doc: Document):
        """Adiciona título principal do documento"""
        # Título principal
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("ESTUDO TÉCNICO PRELIMINAR (ETP)")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Espaço
        doc.add_paragraph()
        
        # Parágrafo introdutório
        intro_para = doc.add_paragraph(
            "O presente documento caracteriza a primeira etapa da fase de planejamento e apresenta "
            "os devidos estudos para a contratação de solução que melhor atenderá à necessidade "
            "descrita abaixo, em observância às normas vigentes e aos princípios que regem a "
            "Administração Pública, especialmente a Lei nº 14.133/2021.",
            style='Corpo Texto ETP'
        )
    
    def _process_and_add_content(self, doc: Document, content: str):
        """Processa e adiciona o conteúdo principal do ETP"""
        lines = content.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Linha vazia - finalizar parágrafo atual se houver
                if current_paragraph:
                    self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                continue
            
            # Identificar tipo de linha
            if self._is_main_section_title(line):
                # Finalizar parágrafo anterior
                if current_paragraph:
                    self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                
                # Adicionar título principal
                self._add_main_section_title(doc, line)
                
            elif self._is_subsection_title(line):
                # Finalizar parágrafo anterior
                if current_paragraph:
                    self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                
                # Adicionar subtítulo
                self._add_subsection_title(doc, line)
                
            elif self._is_table_content(line):
                # Finalizar parágrafo anterior
                if current_paragraph:
                    self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
                    current_paragraph = []
                
                # Processar tabela
                self._add_table_content(doc, line)
                
            else:
                # Conteúdo normal - adicionar ao parágrafo atual
                current_paragraph.append(line)
        
        # Finalizar último parágrafo
        if current_paragraph:
            self._add_paragraph_to_doc(doc, '\n'.join(current_paragraph))
    
    def _is_main_section_title(self, line: str) -> bool:
        """Verifica se a linha é um título de seção principal"""
        pattern = r'^\d+\.\s+[A-ZÁÀÂÃÉÊÍÓÔÕÚÇ\s]+$'
        return bool(re.match(pattern, line.upper()))
    
    def _is_subsection_title(self, line: str) -> bool:
        """Verifica se a linha é um título de subseção"""
        pattern = r'^\d+\.\d+\s+[A-Za-záàâãéêíóôõúç\s]+'
        return bool(re.match(pattern, line))
    
    def _is_table_content(self, line: str) -> bool:
        """Verifica se a linha contém conteúdo de tabela"""
        return '|' in line and line.count('|') >= 2
    
    def _add_main_section_title(self, doc: Document, title: str):
        """Adiciona título de seção principal com fundo azul"""
        para = doc.add_paragraph(title.upper(), style='Titulo Principal ETP')
        
        # Adicionar fundo azul escuro
        self._add_blue_background(para)
    
    def _add_subsection_title(self, doc: Document, title: str):
        """Adiciona título de subseção"""
        doc.add_paragraph(title, style='Subtitulo ETP')
    
    def _add_paragraph_to_doc(self, doc: Document, text: str):
        """Adiciona parágrafo de texto normal"""
        if text.strip():
            doc.add_paragraph(text, style='Corpo Texto ETP')
    
    def _add_table_content(self, doc: Document, table_line: str):
        """Adiciona conteúdo de tabela formatada"""
        # Dividir linha por |
        cells = [cell.strip() for cell in table_line.split('|') if cell.strip()]
        
        if len(cells) >= 2:
            # Criar tabela simples
            table = doc.add_table(rows=1, cols=len(cells))
            table.style = 'Table Grid'
            
            # Adicionar dados
            row = table.rows[0]
            for i, cell_text in enumerate(cells):
                if i < len(row.cells):
                    cell = row.cells[i]
                    cell.text = cell_text
                    
                    # Formatação da célula
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
    
    def _add_blue_background(self, paragraph):
        """Adiciona fundo azul escuro ao parágrafo"""
        try:
            # Criar elemento de sombreamento
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="1f4e79"/>'
            )
            paragraph._element.get_or_add_pPr().append(shading_elm)
        except Exception:
            # Se não conseguir aplicar o fundo, continuar sem ele
            pass
    
    def _add_footer(self, doc: Document):
        """Adiciona rodapé com numeração de páginas"""
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar numeração de página
        run = footer_para.runs[0] if footer_para.runs else footer_para.add_run()
        
        # Criar campo de numeração
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._element.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar2)
        
        # Adicionar texto adicional
        footer_para.add_run(" de ")
        
        # Campo para total de páginas
        run2 = footer_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run2._element.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        run2._element.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run2._element.append(fldChar4)
    
    def _add_signature_section(self, doc: Document, session_data: Dict = None):
        """Adiciona seção de assinatura"""
        # Espaço antes da assinatura
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Linha para assinatura
        sig_para = doc.add_paragraph("_" * 50)
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Nome do responsável
        name_para = doc.add_paragraph("Nome do Responsável pela Elaboração")
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.runs[0]
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(12)
        name_run.font.bold = True
        
        # Cargo
        cargo_para = doc.add_paragraph("Cargo/Função")
        cargo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cargo_run = cargo_para.runs[0]
        cargo_run.font.name = 'Arial'
        cargo_run.font.size = Pt(11)
        
        # Data
        date_para = doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.name = 'Arial'
        date_run.font.size = Pt(11)
    
    def _save_document(self, doc: Document) -> str:
        """Salva o documento e retorna o caminho"""
        temp_dir = tempfile.gettempdir()
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"ETP_Profissional_{timestamp}.docx"
        doc_path = os.path.join(temp_dir, filename)
        
        doc.save(doc_path)
        return doc_path
    
    def create_table_from_data(self, doc: Document, headers: List[str], data: List[List[str]], title: str = None):
        """Cria tabela formatada com dados"""
        if title:
            title_para = doc.add_paragraph(title, style='Subtitulo ETP')
        
        # Criar tabela
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Adicionar cabeçalhos
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            
            # Formatação do cabeçalho
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.bold = True
            
            # Fundo cinza claro para cabeçalho
            try:
                shading_elm = parse_xml(
                    f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="f2f2f2"/>'
                )
                cell._element.get_or_add_tcPr().append(shading_elm)
            except:
                pass
        
        # Adicionar dados
        for row_data in data:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                if i < len(row.cells):
                    cell = row.cells[i]
                    cell.text = str(cell_data)
                    
                    # Formatação da célula
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
        
        # Espaço após tabela
        doc.add_paragraph()

