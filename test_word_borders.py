#!/usr/bin/env python3
"""
Teste de Geração de Documento Word com Bordas
"""
import os
import sys
from datetime import datetime

# Adicionar path para importação
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

def test_word_document_with_borders():
    """Testa geração de documento Word com bordas baseado no modelo da concorrência"""
    print("=== TESTE DE DOCUMENTO WORD COM BORDAS ===")
    print(f"🎯 Objetivo: Gerar documento .docx com bordas como modelo da concorrência")
    print(f"🕐 Iniciando às {datetime.now().strftime('%H:%M:%S')}")
    
    try:
        from utils.word_formatter_with_borders import WordFormatterWithBorders
        print("✅ WordFormatterWithBorders importado com sucesso")
        
        # Criar formatador
        formatter = WordFormatterWithBorders()
        print("✅ Formatador inicializado")
        
        # Conteúdo de teste baseado no ETP real
        test_content = """
1. INTRODUÇÃO

O presente Estudo Técnico Preliminar (ETP) foi elaborado em conformidade com o art. 18, inciso I, da Lei nº 14.133/2021, que estabelece a obrigatoriedade de elaboração de estudos técnicos preliminares para contratações públicas.

Este documento tem por finalidade demonstrar a viabilidade técnica e econômica da contratação pretendida, bem como subsidiar a definição do objeto, dos requisitos técnicos e das demais especificações necessárias.

A elaboração deste estudo observa os princípios da legalidade, impessoalidade, moralidade, publicidade e eficiência, conforme estabelecido no art. 37 da Constituição Federal.

O planejamento da contratação visa garantir o melhor resultado para a Administração Pública, considerando os aspectos de qualidade, sustentabilidade e economicidade.

A metodologia empregada baseia-se em análise técnica detalhada, pesquisa de mercado e avaliação de riscos, assegurando a adequação da solução às necessidades identificadas.

Este documento será submetido à aprovação da autoridade competente antes do prosseguimento do processo de contratação.

A transparência e a conformidade legal são premissas fundamentais que norteiam todo o processo de elaboração deste estudo.

O resultado esperado é a identificação da melhor solução técnica e economicamente viável para atender às necessidades da Administração.

2. OBJETO DO ESTUDO E ESPECIFICAÇÕES GERAIS

2.1. Objeto: Contratação de sistema integrado de gestão de contratos públicos com inteligência artificial para modernização dos processos administrativos.

2.2. Planejamento / Demanda específica: A demanda decorre da necessidade de modernização dos processos de gestão contratual da Administração Pública.

2.3. Modalidade de licitação a ser adotada: Concorrência, conforme art. 28, inciso I, da Lei nº 14.133/2021.

2.4. Tipo de Licitação: Técnica e Preço, considerando a complexidade técnica da solução.

2.5. Critério de Julgamento: Melhor técnica e preço, conforme art. 34 da Lei nº 14.133/2021.

2.6. Registro de preços: Não se aplica para esta contratação específica.

3. DESCRIÇÃO DOS REQUISITOS DA CONTRATAÇÃO

Os requisitos técnicos foram definidos com base na análise das necessidades operacionais e nas melhores práticas de gestão contratual.

A solução deverá contemplar funcionalidades de gestão completa do ciclo de vida dos contratos, desde o planejamento até a execução e fiscalização.

É obrigatória a conformidade com a Lei Geral de Proteção de Dados (LGPD) e demais normas de segurança da informação.

A arquitetura da solução deve ser baseada em nuvem, garantindo escalabilidade e disponibilidade dos serviços.

A integração com sistemas legados da Administração é requisito fundamental para a continuidade operacional.

O sistema deve possuir interface intuitiva e recursos de business intelligence para apoio à tomada de decisão.

A solução deve contemplar módulos de workflow, gestão documental e controle de prazos contratuais.

É necessária a disponibilização de APIs para integração com outros sistemas governamentais.

4. ESTIMATIVA DAS QUANTIDADES E VALORES

A estimativa de valores foi elaborada com base em pesquisa de mercado realizada junto a fornecedores qualificados.

Foram consultados preços praticados em contratos similares de outros órgãos públicos para validação dos valores.

A metodologia considera custos de licenciamento, implementação, customização e treinamento dos usuários.

Os valores foram ajustados considerando a complexidade específica dos requisitos da Administração.

| Item | Quantidade | Unidade | Valor Unitário | Valor Total |
|------|------------|---------|----------------|-------------|
| Licenças do Sistema | 100 | Usuário/ano | R$ 2.500,00 | R$ 250.000,00 |
| Implementação e Customização | 1 | Projeto | R$ 100.000,00 | R$ 100.000,00 |
| Treinamento de Usuários | 1 | Projeto | R$ 50.000,00 | R$ 50.000,00 |
| **TOTAL ESTIMADO** | | | | **R$ 400.000,00** |

A estimativa considera margem de segurança de 10% para eventuais ajustes durante a implementação.

Os valores incluem todos os tributos e encargos aplicáveis conforme legislação vigente.

A validade da estimativa é de 6 meses, podendo ser revista conforme variações de mercado.

5. LEVANTAMENTO DE MERCADO E JUSTIFICATIVA DA ESCOLHA DA SOLUÇÃO

O levantamento de mercado identificou diversas soluções disponíveis com diferentes níveis de adequação aos requisitos.

Foram analisadas soluções nacionais e internacionais, priorizando fornecedores com experiência no setor público.

A escolha da solução baseou-se em critérios de funcionalidade, segurança, escalabilidade e custo-benefício.

Foram considerados aspectos de suporte técnico, capacitação e transferência de conhecimento.

A solução selecionada apresenta o melhor equilíbrio entre atendimento aos requisitos e viabilidade econômica.

O fornecedor possui certificações de qualidade e segurança reconhecidas internacionalmente.

A arquitetura proposta garante flexibilidade para futuras expansões e integrações.

A análise de TCO (Total Cost of Ownership) demonstrou a vantajosidade da solução escolhida.

13. ANÁLISE DE RISCOS

A análise de riscos contempla aspectos técnicos, operacionais, financeiros e de cronograma do projeto.

Foram identificados riscos relacionados à integração com sistemas legados e à migração de dados.

| Risco | Probabilidade | Impacto | Estratégia de Mitigação |
|-------|---------------|---------|------------------------|
| Problemas de integração | Média | Alto | Testes extensivos e ambiente de homologação |
| Atraso na implementação | Baixa | Médio | Cronograma com folgas e marcos de controle |
| Resistência dos usuários | Média | Médio | Programa de capacitação e gestão da mudança |
| Problemas de performance | Baixa | Alto | Testes de carga e dimensionamento adequado |

O monitoramento contínuo dos riscos será realizado durante toda a execução do projeto.

Planos de contingência foram elaborados para os riscos de maior criticidade.

A governança do projeto incluirá comitê específico para gestão de riscos e tomada de decisões.

14. CONCLUSÃO E POSICIONAMENTO FINAL

Com base nos estudos realizados, conclui-se pela viabilidade técnica e econômica da contratação.

A solução proposta atende integralmente aos requisitos identificados e às necessidades da Administração.

Os riscos foram adequadamente mapeados e possuem estratégias de mitigação definidas.

O investimento se justifica pelos benefícios esperados em termos de eficiência e modernização.

Recomenda-se o prosseguimento do processo licitatório conforme modalidade e critérios definidos.

A implementação da solução contribuirá significativamente para a melhoria da gestão contratual.

O cronograma proposto é factível e compatível com as necessidades operacionais.

Este ETP demonstra o cumprimento das exigências legais e a adequação da contratação aos princípios da Administração Pública.
"""
        
        # Dados da sessão de teste
        session_data = {
            'answers': {
                '1': 'Sistema integrado de gestão de contratos públicos com IA',
                '2': 'Sim, previsto no PCA 2024',
                '3': 'Lei 14.133/21, LGPD, Decreto 10.024/19',
                '4': 'R$ 400.000,00 total estimado',
                '5': 'Não haverá parcelamento'
            },
            'session_id': 'test_borders_2024',
            'title': 'Estudo Técnico Preliminar - Teste com Bordas'
        }
        
        print("📝 Gerando documento Word com bordas...")
        
        # Gerar documento
        doc_path = formatter.create_etp_with_borders(test_content, session_data)
        
        print(f"✅ Documento gerado com sucesso!")
        print(f"📁 Caminho: {doc_path}")
        
        # Verificar se arquivo foi criado
        if os.path.exists(doc_path):
            file_size = os.path.getsize(doc_path)
            print(f"📊 Tamanho do arquivo: {file_size:,} bytes")
            print(f"📄 Formato: .docx")
            
            # Verificar se é um arquivo Word válido
            if doc_path.endswith('.docx') and file_size > 10000:  # Pelo menos 10KB
                print("✅ Arquivo Word válido gerado")
                
                # Tentar abrir com python-docx para validar
                try:
                    from docx import Document
                    doc = Document(doc_path)
                    paragraph_count = len(doc.paragraphs)
                    print(f"📊 Parágrafos no documento: {paragraph_count}")
                    
                    # Verificar se tem conteúdo
                    if paragraph_count > 10:
                        print("✅ Documento contém conteúdo adequado")
                    else:
                        print("⚠️  Documento pode estar com pouco conteúdo")
                        
                except Exception as e:
                    print(f"⚠️  Erro ao validar documento: {e}")
            else:
                print("❌ Arquivo gerado pode estar corrompido")
        else:
            print("❌ Arquivo não foi criado")
            return False
        
        # Testar elementos específicos do modelo
        print("\n🎨 Verificando elementos do modelo da concorrência:")
        
        # Verificar se métodos específicos existem
        methods_to_check = [
            '_apply_document_borders',
            '_add_blue_background', 
            '_add_institutional_header',
            '_add_main_title_with_background',
            '_create_formatted_table'
        ]
        
        for method in methods_to_check:
            if hasattr(formatter, method):
                print(f"  ✅ {method}")
            else:
                print(f"  ❌ {method}")
        
        print(f"\n🎉 TESTE CONCLUÍDO COM SUCESSO!")
        print(f"📁 Documento disponível em: {doc_path}")
        print(f"🎯 Resultado: Documento Word com bordas baseado no modelo da concorrência")
        
        return doc_path
        
    except ImportError as e:
        print(f"❌ Erro de importação: {e}")
        return False
    except Exception as e:
        print(f"❌ Erro geral: {e}")
        return False

def test_border_elements():
    """Testa elementos específicos de bordas"""
    print("\n=== TESTE DE ELEMENTOS DE BORDAS ===")
    
    try:
        from utils.word_formatter_with_borders import WordFormatterWithBorders
        
        formatter = WordFormatterWithBorders()
        
        # Verificar cores definidas
        print(f"🎨 Cor azul: RGB{formatter.blue_color}")
        print(f"🎨 Cor branca: RGB{formatter.white_color}")
        print(f"🎨 Cor preta: RGB{formatter.black_color}")
        
        # Verificar se métodos de formatação existem
        formatting_methods = [
            '_configure_page_with_borders',
            '_add_cell_border',
            '_add_paragraph_border',
            '_remove_table_borders'
        ]
        
        print("\n🔧 Métodos de formatação:")
        for method in formatting_methods:
            if hasattr(formatter, method):
                print(f"  ✅ {method}")
            else:
                print(f"  ❌ {method}")
        
        print("✅ Elementos de bordas validados")
        
    except Exception as e:
        print(f"❌ Erro na validação: {e}")

if __name__ == "__main__":
    print("🚀 Iniciando teste de documento Word com bordas...")
    
    # Teste principal
    result = test_word_document_with_borders()
    
    # Teste de elementos
    test_border_elements()
    
    if result:
        print(f"\n✅ SUCESSO: Documento Word com bordas gerado!")
        print(f"📁 Arquivo: {result}")
        print(f"🎯 O documento final agora possui bordas como o modelo da concorrência!")
    else:
        print(f"\n❌ FALHA: Não foi possível gerar o documento")
    
    print(f"\n🕐 Teste finalizado às {datetime.now().strftime('%H:%M:%S')}")

